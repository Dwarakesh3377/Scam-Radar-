"""
Generate Algorithms Word Document - UG Student Academic Format
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color_hex):
    shd_elem = OxmlElement('w:shd')
    shd_elem.set(qn('w:fill'), color_hex)
    shd_elem.set(qn('w:val'), 'clear')
    cell._element.get_or_add_tcPr().append(shd_elem)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, 'D9E2EC')
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(value))
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(width)
    return table

def add_caption(doc, number, text):
    p = doc.add_paragraph('')
    run = p.add_run(f'Table {number}: ')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    r2.font.name = 'Times New Roman'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_algorithm_box(doc, algo_number, algo_name, input_line, output_line, steps):
    """Add algorithm in standard UG report format with border"""
    # Create a 1-cell table to act as a bordered box
    box = doc.add_table(rows=1, cols=1)
    box.style = 'Table Grid'
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = box.rows[0].cells[0]
    cell.text = ''
    set_cell_shading(cell, 'F8F9FA')

    # Algorithm title
    p = cell.paragraphs[0]
    run = p.add_run(f'Algorithm {algo_number}: {algo_name}')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(6)

    # Input
    p = cell.add_paragraph()
    run = p.add_run('Input: ')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    r = p.add_run(input_line)
    r.font.size = Pt(11)
    r.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)

    # Output
    p = cell.add_paragraph()
    run = p.add_run('Output: ')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    r = p.add_run(output_line)
    r.font.size = Pt(11)
    r.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(2)

    # BEGIN
    p = cell.add_paragraph()
    run = p.add_run('BEGIN')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(4)

    # Steps
    for step_text in steps:
        p = cell.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.line_spacing = 1.15

        if step_text == '':
            # blank spacer
            run = p.add_run(' ')
            run.font.size = Pt(6)
            run.font.name = 'Times New Roman'
        elif step_text.strip().startswith('Step'):
            # "Step N:" bold, rest normal
            colon_idx = step_text.find(':')
            if colon_idx > 0:
                run = p.add_run(step_text[:colon_idx + 1])
                run.bold = True
                run.font.size = Pt(11)
                run.font.name = 'Times New Roman'
                r = p.add_run(step_text[colon_idx + 1:])
                r.font.size = Pt(11)
                r.font.name = 'Times New Roman'
            else:
                run = p.add_run(step_text)
                run.font.size = Pt(11)
                run.font.name = 'Times New Roman'
        else:
            # Continuation / indented line
            p.paragraph_format.left_indent = Cm(1.5)
            run = p.add_run(step_text)
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'

    # END
    p = cell.add_paragraph()
    run = p.add_run('END')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_before = Pt(4)

    doc.add_paragraph('')  # spacing after box


def create_document():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(2.54)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)

    for level in range(1, 4):
        h = doc.styles[f'Heading {level}']
        h.font.name = 'Times New Roman'
        h.font.color.rgb = RGBColor(0, 0, 0)

    # ══════════════════════════════════════════════════
    # TITLE PAGE
    # ══════════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph('')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ALGORITHMS USED IN THE PROJECT')
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = 'Times New Roman'

    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Scam Radar \u2014 AI-Powered Scam Risk Detection System')
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # INTRODUCTION
    # ══════════════════════════════════════════════════
    doc.add_heading('Introduction', level=1)

    doc.add_paragraph(
        'The Scam Radar application uses a combination of Deep Learning, Machine Learning, '
        'and Rule-Based algorithms to detect whether a given text input (such as a job posting, '
        'email, URL, or phone number) is legitimate, suspicious, or a scam. The system supports '
        '10 languages and produces a risk score from 0 to 100.'
    )
    doc.add_paragraph(
        'The following six algorithms are the core of the prediction engine. Three of them '
        '(BERT, TF-IDF + Random Forest, and Rule-Based Scoring) run independently and produce '
        'their own risk scores. These three scores are then combined by the Weighted Ensemble '
        'algorithm to produce the final result.'
    )

    doc.add_paragraph('')
    add_table(doc,
        ['Algorithm', 'Type', 'Ensemble Weight'],
        [
            ['BERT / XLM-RoBERTa', 'Deep Learning', '30%'],
            ['TF-IDF + Random Forest', 'Machine Learning', '20%'],
            ['Rule-Based Scoring', 'Heuristic / Expert System', '50%'],
        ],
        col_widths=[5.5, 4.5, 4]
    )
    add_caption(doc, 1, 'Ensemble weight distribution.')

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # ALGORITHM 1 — BERT
    # ══════════════════════════════════════════════════
    doc.add_heading('Algorithm 1: BERT Prediction', level=1)

    doc.add_paragraph(
        'BERT (Bidirectional Encoder Representations from Transformers) is a deep learning model '
        'that reads text in both directions to understand context. It is fine-tuned to classify '
        'English text into three classes: Legitimate, Suspicious, and Scam.'
    )

    add_algorithm_box(doc, 1, 'BERT English Text Classification',
        'text (the English text to be analyzed)',
        'risk_score (0 to 100), confidence (0.0 to 1.0)',
        [
            'Step 1: Check if the BERT model is loaded in memory.',
            'If the model is not loaded, return score = 0 and confidence = 0.',
            '',
            'Step 2: Tokenize the input text using the BERT tokenizer.',
            'Convert the text into numerical tokens that the model can process.',
            'Set maximum token length to 512.',
            '',
            'Step 3: Pass the tokens through the BERT model to get raw output values (logits).',
            '',
            'Step 4: Apply the Softmax function on the logits to get three probability values:',
            'P(Legitimate), P(Suspicious), P(Scam).',
            'These three values will sum to 1.0.',
            '',
            'Step 5: Calculate the risk score using the formula:',
            'risk_score = (P(Scam) \u00d7 100) + (P(Suspicious) \u00d7 50)',
            '',
            'Step 6: Calculate confidence as the highest probability among the three classes:',
            'confidence = max(P(Legitimate), P(Suspicious), P(Scam))',
            '',
            'Step 7: If risk_score exceeds 100, set it to 100.',
            '',
            'Step 8: Return the risk_score and confidence.',
        ]
    )

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Source File: ')
    run.bold = True
    p.add_run('backend/services/predict.py \u2192 get_bert_prediction()')

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # ALGORITHM 2 — XLM-RoBERTa
    # ══════════════════════════════════════════════════
    doc.add_heading('Algorithm 2: XLM-RoBERTa Multilingual Prediction', level=1)

    doc.add_paragraph(
        'XLM-RoBERTa is a multilingual deep learning model that supports 100 languages. '
        'It is used when the input text is in a non-English language such as Tamil, Hindi, '
        'French, Spanish, German, Russian, Chinese, Japanese, or Korean.'
    )

    add_algorithm_box(doc, 2, 'XLM-RoBERTa Multilingual Text Classification',
        'text (non-English text to be analyzed)',
        'risk_score (0 to 100), confidence (0.0 to 1.0)',
        [
            'Step 1: Check if the XLM-RoBERTa model is loaded in memory.',
            'If the model is not loaded, return score = 0 and confidence = 0.',
            '',
            'Step 2: Tokenize the input text using the XLM-RoBERTa tokenizer.',
            'Set maximum token length to 512.',
            '',
            'Step 3: Pass the tokens through the XLM-RoBERTa model to get logits.',
            '',
            'Step 4: Apply the Softmax function to get three probabilities:',
            'P(Legitimate), P(Suspicious), P(Scam).',
            '',
            'Step 5: Calculate the risk score:',
            'risk_score = (P(Scam) \u00d7 100) + (P(Suspicious) \u00d7 50)',
            '',
            'Step 6: Calculate the confidence:',
            'confidence = max(P(Legitimate), P(Suspicious), P(Scam))',
            '',
            'Step 7: If risk_score exceeds 100, set it to 100.',
            '',
            'Step 8: Return the risk_score and confidence.',
        ]
    )

    doc.add_paragraph('')
    doc.add_paragraph(
        'Note: If XLM-RoBERTa fails to produce a valid result (score = 0 and confidence = 0), '
        'the system automatically falls back to the English BERT model as a safety measure.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Source File: ')
    run.bold = True
    p.add_run('backend/services/predict.py \u2192 get_xlm_prediction()')

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # ALGORITHM 3 — TF-IDF + RANDOM FOREST
    # ══════════════════════════════════════════════════
    doc.add_heading('Algorithm 3: TF-IDF with Random Forest Classifier', level=1)

    doc.add_paragraph(
        'TF-IDF (Term Frequency\u2013Inverse Document Frequency) is a technique that converts '
        'text into numbers by calculating how important each word is. Random Forest is a '
        'machine learning algorithm that combines many decision trees and uses majority voting '
        'to make predictions. Together, they act as a lightweight baseline model.'
    )

    add_algorithm_box(doc, 3, 'TF-IDF + Random Forest Prediction',
        'text (the text to be analyzed)',
        'risk_score (0 to 100), confidence (0.0 to 1.0)',
        [
            'Step 1: Check if the TF-IDF model and vectorizer are loaded.',
            'If not loaded, return score = NULL and confidence = 0.5.',
            '',
            'Step 2: Convert the input text into a TF-IDF feature vector.',
            'The vectorizer calculates the importance score for each word.',
            '',
            'Step 3: If a scaler is available, normalize the feature values.',
            '',
            'Step 4: Feed the feature vector to the Random Forest model.',
            'The model outputs the predicted class and probability for each class.',
            '',
            'Step 5: Map the predicted class to a risk score range:',
            'If predicted class = 0 (Legitimate):',
            '    risk_score = 10 + (1 \u2212 P(Legitimate)) \u00d7 23    [Range: 10 to 33]',
            'If predicted class = 1 (Suspicious):',
            '    risk_score = 40 + P(Scam) \u00d7 26              [Range: 40 to 66]',
            'If predicted class = 2 (Scam):',
            '    risk_score = 70 + P(Scam) \u00d7 30              [Range: 70 to 100]',
            '',
            'Step 6: Calculate confidence as the highest probability value.',
            '',
            'Step 7: Return the risk_score and confidence.',
        ]
    )

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Source File: ')
    run.bold = True
    p.add_run('backend/services/predict.py \u2192 predict_with_tfidf()')

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # ALGORITHM 4 — RULE-BASED SCORING
    # ══════════════════════════════════════════════════
    doc.add_heading('Algorithm 4: Rule-Based Heuristic Scoring', level=1)

    doc.add_paragraph(
        'The Rule-Based Scoring algorithm uses handcrafted domain knowledge to assign a '
        'risk score. It checks for known scam patterns (like payment demands, urgency words) '
        'and legitimate patterns (like corporate emails, professional keywords). It carries '
        'the highest weight (50%) in the ensemble because payment and financial fraud patterns '
        'must always be caught, regardless of what the AI models predict.'
    )

    add_algorithm_box(doc, 4, 'Rule-Based Heuristic Risk Scoring',
        'features (the extracted feature values from the input text)',
        'score (0 to 100)',
        [
            'Step 1: Start with a base score of 15.',
            '(A low base score means the text is assumed to be legitimate initially.)',
            '',
            'Step 2: Check for financial scam keywords (e.g., "payment", "deposit", "transfer").',
            'If found, add 100 points to the score.',
            '',
            'Step 3: Check for urgency keywords (e.g., "urgent", "immediate", "act now").',
            'If more than 1 urgency word is found, add 25 points.',
            'If exactly 1 urgency word is found, add 10 points.',
            '',
            'Step 4: Check for other scam indicators:',
            'If WhatsApp or Telegram is mentioned, add 20 points.',
            'If sender uses a free email (Gmail, Yahoo), add 20 points.',
            'If domain has suspicious extension (.xyz, .tk), add 40 points.',
            'If domain name mismatches the claimed company, add 30 points.',
            '',
            'Step 5: Check for corporate email (non-free email domain).',
            'If the sender has a corporate email, subtract 25 points.',
            '',
            'Step 6: Check for professional/legitimate keywords',
            '(e.g., "interview", "experience required", "job description").',
            'If found, subtract 10 to 35 points based on count.',
            '',
            'Step 7: Check for HTTPS in the URL.',
            'If present, subtract 10 points.',
            '',
            'Step 8: If zero scam keywords were found, subtract 15 points.',
            '',
            'Step 9: Clamp the final score between 0 and 100.',
            'If score < 0, set score = 0.',
            'If score > 100, set score = 100.',
            '',
            'Step 10: Return the score.',
        ]
    )

    doc.add_paragraph('')
    add_table(doc,
        ['Indicator Type', 'Feature', 'Points'],
        [
            ['Scam', 'Financial keywords (payment, deposit)', '+100'],
            ['Scam', 'Suspicious domain (.xyz, .tk)', '+40'],
            ['Scam', 'Domain mismatch', '+30'],
            ['Scam', 'Multiple urgency keywords', '+25'],
            ['Scam', 'Free email provider (Gmail, Yahoo)', '+20'],
            ['Scam', 'WhatsApp / Telegram contact', '+20'],
            ['Scam', 'Single urgency keyword', '+10'],
            ['Legitimate', 'Corporate email sender', '\u221225'],
            ['Legitimate', 'Professional keywords', '\u221210 to \u221235'],
            ['Legitimate', 'No scam keywords found', '\u221215'],
            ['Legitimate', 'HTTPS URL', '\u221210'],
        ],
        col_widths=[3, 6, 3]
    )
    add_caption(doc, 2, 'Rule-based scoring weight table.')

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Source File: ')
    run.bold = True
    p.add_run('backend/services/predict.py \u2192 calculate_rule_based_score()')

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # ALGORITHM 5 — LANGUAGE DETECTION
    # ══════════════════════════════════════════════════
    doc.add_heading('Algorithm 5: Language Detection and Model Routing', level=1)

    doc.add_paragraph(
        'This algorithm detects what language the user has typed in and then selects the '
        'correct AI model. English text goes to BERT. All other languages go to XLM-RoBERTa.'
    )

    add_algorithm_box(doc, 5, 'Language Detection and Model Routing',
        'text (the raw text entered by the user)',
        'language_code, selected_model, score, confidence',
        [
            'Step 1: If the text is empty or has fewer than 10 characters,',
            'assume the language is English and select BERT model.',
            '',
            'Step 2: Use the langdetect library to detect the language of the text.',
            '',
            'Step 3: Check if the detected language is in the supported list:',
            'Supported languages: English, Tamil, Hindi, French, Spanish,',
            'German, Japanese, Chinese, Russian, Korean.',
            '',
            'Step 4: If the detected language is not supported,',
            'default to English.',
            '',
            'Step 5: Select the AI model based on language:',
            'If language is English \u2192 select BERT model.',
            'If language is any other supported language \u2192 select XLM-RoBERTa.',
            '',
            'Step 6: Run the selected model on the text to get the prediction.',
            '',
            'Step 7: If XLM-RoBERTa returns an empty result (score = 0, confidence = 0),',
            'fall back to BERT and run the prediction again.',
            '',
            'Step 8: Return the language code, selected model, score, and confidence.',
        ]
    )

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Source File: ')
    run.bold = True
    p.add_run('backend/services/language.py \u2192 detect_language()')

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # ALGORITHM 6 — WEIGHTED ENSEMBLE
    # ══════════════════════════════════════════════════
    doc.add_heading('Algorithm 6: Weighted Ensemble Risk Prediction', level=1)

    doc.add_paragraph(
        'This is the main algorithm of the project. It takes the scores from all three '
        'models (BERT/XLM-RoBERTa, TF-IDF, and Rule-Based), combines them using fixed weights, '
        'and then applies several safety checks to produce the final risk score and risk level.'
    )

    add_algorithm_box(doc, 6, 'Weighted Ensemble Risk Prediction',
        'text, metadata (sender email, phone, etc.), language',
        'final_score (0 to 100), confidence (0.35 to 0.95), risk_level',
        [
            'Step 1: Load all AI models into memory if not already loaded.',
            '',
            'Step 2: Detect the language of the text.',
            'If English, get prediction from BERT model.',
            'If other language, get prediction from XLM-RoBERTa.',
            'If XLM-RoBERTa fails, fall back to BERT.',
            'Store the result as ml_score and ml_confidence.',
            '',
            'Step 3: Get prediction from TF-IDF + Random Forest model.',
            'Store the result as tfidf_score and tfidf_confidence.',
            '',
            'Step 4: Get score from the Rule-Based algorithm.',
            'Store the result as rule_score.',
            '',
            'Step 5: Combine the three scores using weights:',
            'final_score = (rule_score \u00d7 0.5) + (ml_score \u00d7 0.3) + (tfidf_score \u00d7 0.2)',
            'confidence = maximum of (ml_confidence, tfidf_confidence, 0.40)',
            '',
            'Step 6: SHORT TEXT PROTECTION \u2014',
            'If the input text is very short (< 25 characters) or mostly numbers,',
            'and no financial keywords are found,',
            'cap the score at 40 to avoid false alarms.',
            '',
            'Step 7: SCAM BOOSTER \u2014',
            'If financial keywords are found (e.g., "payment", "deposit"),',
            'boost the score to at least 85.',
            '',
            'Step 8: RED FLAG COUNTER \u2014',
            'Count the number of red flags (suspicious TLD, urgency words,',
            'domain mismatch, free email, scam contact methods).',
            'If 3 or more red flags: set minimum score to 70.',
            'If 2 red flags: set minimum score to 55.',
            'If 1 red flag: set minimum score to 40.',
            '',
            'Step 9: LEGITIMATE PROTECTOR \u2014',
            'If there are zero scam keywords, no financial terms,',
            'no free email, no suspicious TLD, and no domain mismatch,',
            'cap the score at 20 (clearly safe).',
            '',
            'Step 10: CORPORATE PROTECTOR \u2014',
            'If the sender uses a corporate email,',
            'and there are no financial or urgency keywords,',
            'cap the score at 20 (if legitimate keywords present) or 30.',
            '',
            'Step 11: FINAL OVERRIDE \u2014',
            'If financial keywords are present (checked again after protectors),',
            'force the score to at least 80.',
            'This ensures payment demands are never marked as safe.',
            '',
            'Step 12: Clamp the final score to the range 0 to 100.',
            'Add a small random variation (\u00b10.10) to the confidence.',
            'Clamp the confidence to the range 0.35 to 0.95.',
            '',
            'Step 13: Classify the risk level:',
            'Score 0 to 30    \u2192  LEGITIMATE (safe)',
            'Score 31 to 60   \u2192  SUSPICIOUS (caution)',
            'Score 61 to 100  \u2192  SCAM (high risk)',
            '',
            'Step 14: Return the final_score, confidence, and risk_level.',
        ]
    )

    doc.add_paragraph('')
    add_table(doc,
        ['Score Range', 'Risk Level', 'Meaning'],
        [
            ['0 \u2013 30', 'LEGITIMATE', 'The content appears safe and professional'],
            ['31 \u2013 60', 'SUSPICIOUS', 'Some warning signs \u2014 verify before proceeding'],
            ['61 \u2013 100', 'SCAM', 'High risk \u2014 likely fraudulent content'],
        ],
        col_widths=[3, 4, 8]
    )
    add_caption(doc, 3, 'Risk level classification.')

    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run('Source File: ')
    run.bold = True
    p.add_run('backend/services/predict.py \u2192 predict_risk()')

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # SUMMARY
    # ══════════════════════════════════════════════════
    doc.add_heading('Summary of Algorithms', level=1)

    add_table(doc,
        ['No.', 'Algorithm Name', 'Type', 'Purpose'],
        [
            ['1', 'BERT', 'Deep Learning', 'Classify English text into Legitimate / Suspicious / Scam'],
            ['2', 'XLM-RoBERTa', 'Deep Learning', 'Classify non-English text (Tamil, Hindi, French, etc.)'],
            ['3', 'TF-IDF + Random Forest', 'Machine Learning', 'Lightweight baseline using word frequency patterns'],
            ['4', 'Rule-Based Scoring', 'Heuristic', 'Score based on known scam and legitimate patterns'],
            ['5', 'Language Detection', 'NLP', 'Detect input language and route to correct model'],
            ['6', 'Weighted Ensemble', 'Ensemble', 'Combine all scores and apply safety overrides'],
        ],
        col_widths=[1, 4.5, 3, 6.5]
    )
    add_caption(doc, 4, 'Summary of all algorithms used in Scam Radar.')

    # ─── Save ───
    output_path = r'd:\scam-risk-detection2\Algorithms_Used_In_Project.docx'
    doc.save(output_path)
    print(f"\u2705 Document saved: {output_path}")


if __name__ == '__main__':
    create_document()
