from docx import Document
import re

def update_chapter7_flow(file_path):
    doc = Document(file_path)
    
    # Target: Step 1 description which mentions two tabs
    old_header = 'Step 1: User Selects Input Type and Pastes Content'
    old_body_start = 'The Home page presents two tabs'
    
    new_header = "Step 1: User Selects Platform Source and Enters Content"
    new_body = (
        "The Home page features a unified analysis form with a dynamic **Source Platform** dropdown "
        "(LinkedIn, Naukri, Gmail, etc.). When a user selects a platform, the system automatically "
        "adjusts the visible input fields—such as **Sender Email** for email sources or **Company Website** "
        "for job boards. The user pastes their content into the single central textarea, and the system "
        "intelligently determines the input type (job or email) based on the selected platform, "
        "ensuring a seamless and intuitive data entry experience."
    )
    
    updated = False
    for i, p in enumerate(doc.paragraphs):
        if old_header in p.text:
            p.text = new_header
            # The next paragraph should be the body
            if i + 1 < len(doc.paragraphs) and old_body_start in doc.paragraphs[i+1].text:
                doc.paragraphs[i+1].text = new_body
                updated = True
                print(f"Successfully updated 'Home Tab' section in {file_path}")
                break
            
    if not updated:
        print("Warning: Could not find the exact 'Home Tab' section to update.")
    else:
        doc.save(file_path)

if __name__ == "__main__":
    update_chapter7_flow('d:\\scam-risk-detection2\\Chapter7 EndToEnd Flow.docx')
